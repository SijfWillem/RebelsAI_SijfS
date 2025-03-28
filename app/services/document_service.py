import os
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, List, Tuple, Any
import asyncio
import aiofiles
from app.models.document import Document, DocumentType, DocumentStatus, FolderAnalysis, SentimentAnalysis
from app.services.classification_service import ClassificationService
from functools import lru_cache
import time
import logging
from textblob import TextBlob
import magic
import hashlib
from docx import Document as DocxDocument
from app.services.database_service import DatabaseService
from app.database import AsyncSessionLocal

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentService:
    def __init__(self, base_directory: str = "Client Data", batch_size: int = 50):
        self.base_directory = Path(base_directory)
        self.classification_service = ClassificationService()
        self.batch_size = batch_size
        
        # Create base directory if it doesn't exist
        self.base_directory.mkdir(parents=True, exist_ok=True)
        logger.info(f"Initialized DocumentService with base directory: {self.base_directory.absolute()}")

    @lru_cache(maxsize=100)
    def _get_file_type(self, file_path: str) -> DocumentType:
        """Determine the file type based on extension and content with caching."""
        try:
            # First try to get the MIME type using python-magic
            mime = magic.Magic(mime=True)
            mime_type = mime.from_file(str(file_path))
            
            # Map MIME types to document types
            mime_map = {
                'application/pdf': DocumentType.PDF,
                'application/msword': DocumentType.DOCX,
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocumentType.DOCX,
                'text/plain': DocumentType.TXT,
                'application/vnd.ms-excel': DocumentType.XLSX,
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': DocumentType.XLSX,
                'application/vnd.ms-powerpoint': DocumentType.PPTX,
                'application/vnd.openxmlformats-officedocument.presentationml.presentation': DocumentType.PPTX,
                'image/jpeg': DocumentType.JPG,
                'image/png': DocumentType.PNG
            }
            
            doc_type = mime_map.get(mime_type, DocumentType.OTHER)
            logger.info(f"Determined file type for {file_path}: {doc_type} (MIME: {mime_type})")
            return doc_type
        except Exception as e:
            logger.error(f"Error determining file type for {file_path}: {str(e)}")
            # Fallback to extension-based detection
            ext = Path(file_path).suffix.lower().lstrip('.')
            extension_map = {
                'pdf': DocumentType.PDF,
                'doc': DocumentType.DOCX,
                'docx': DocumentType.DOCX,
                'txt': DocumentType.TXT,
                'xls': DocumentType.XLSX,
                'xlsx': DocumentType.XLSX,
                'ppt': DocumentType.PPTX,
                'pptx': DocumentType.PPTX,
                'jpg': DocumentType.JPG,
                'jpeg': DocumentType.JPG,
                'png': DocumentType.PNG
            }
            return extension_map.get(ext, DocumentType.OTHER)

    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of a file."""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()

    async def extract_text_content(self, file_path: Path) -> Optional[str]:
        """Extract text content from various file types."""
        try:
            file_type = self._get_file_type(str(file_path))
            
            if file_type == DocumentType.DOCX:
                doc = DocxDocument(file_path)
                return "\n".join([paragraph.text for paragraph in doc.paragraphs])
            elif file_type == DocumentType.TXT:
                async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                    return await f.read()
            elif file_type == DocumentType.PDF:
                # TODO: Implement PDF text extraction
                return None
            else:
                return None
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None

    async def analyze_sentiment(self, text: str) -> Dict[str, Any]:
        """Analyze sentiment of text using TextBlob for multilingual support."""
        try:
            # Use TextBlob for sentiment analysis
            blob = TextBlob(text)
            
            # Get polarity (-1 to 1) and subjectivity (0 to 1)
            polarity = blob.sentiment.polarity
            subjectivity = blob.sentiment.subjectivity
            
            # Determine sentiment category
            if polarity > 0.1:
                sentiment_category = "positive"
            elif polarity < -0.1:
                sentiment_category = "negative"
            else:
                sentiment_category = "neutral"
            
            return {
                "sentiment": sentiment_category,
                "polarity": polarity,
                "subjectivity": subjectivity,
                "confidence": abs(polarity)  # Use absolute polarity as confidence
            }
        except Exception as e:
            logger.error(f"Error analyzing sentiment: {str(e)}")
            return {
                "sentiment": "neutral",
                "polarity": 0,
                "subjectivity": 0,
                "confidence": 0
            }

    async def create_document(self, file_path: Path, base_dir: Path) -> Document:
        """Create a Document object from a file path."""
        try:
            # Get relative path from base directory
            relative_path = file_path.relative_to(base_dir)
            logger.info(f"Creating document from file: {relative_path}")

            # Get file type
            file_type = self._get_file_type(str(file_path))
            logger.info(f"File type detected: {file_type}")

            # Get file stats
            stats = file_path.stat()
            size = stats.st_size
            created_at = datetime.fromtimestamp(stats.st_ctime)
            modified_at = datetime.fromtimestamp(stats.st_mtime)
            logger.info(f"File stats - Size: {size}, Created: {created_at}, Modified: {modified_at}")

            # Process text content for analysis
            content = None
            sentiment = None
            
            if file_type in [DocumentType.TXT, DocumentType.MARKDOWN, DocumentType.PDF, DocumentType.DOCX]:
                try:
                    content = await self.extract_text_content(file_path)
                    if content:
                        # Perform sentiment analysis
                        sentiment = await self.analyze_sentiment(content)
                        logger.info(f"Analysis completed for {relative_path}")
                except Exception as e:
                    logger.error(f"Error processing content for {relative_path}: {str(e)}")

            # Create document object with all required fields
            doc = Document(
                id=str(uuid.uuid4()),
                filename=str(relative_path),
                file_type=file_type,
                size=size,
                created_at=created_at,
                modified_at=modified_at,
                path=str(relative_path),
                content=content,
                sentiment=sentiment,
                metadata={
                    "mime_type": magic.Magic(mime=True).from_file(str(file_path)),
                    "processed": True
                }
            )

            # Save to database
            async with AsyncSessionLocal() as session:
                db_service = DatabaseService(session)
                await db_service.create_document(doc)

            logger.info(f"Document created successfully: {doc.filename}")
            return doc

        except Exception as e:
            logger.error(f"Error creating document from {file_path}: {str(e)}", exc_info=True)
            raise

    async def analyze_folder(self, folder_path: Optional[str] = None, page: int = 1, page_size: int = 50) -> Tuple[FolderAnalysis, int]:
        """Analyze contents of a folder with enhanced insights and pagination."""
        try:
            # Get the target directory
            target_dir = Path(folder_path) if folder_path else self.base_directory
            logger.info(f"Analyzing folder. Target directory: {target_dir.absolute()}")
            
            if not target_dir.exists():
                logger.error(f"Target directory does not exist: {target_dir.absolute()}")
                return FolderAnalysis(
                    total_documents=0,
                    total_size=0,
                    document_types={},
                    average_file_size=0,
                    last_modified=datetime.now(),
                    documents=[]
                ), 0

            # Create root folder in database
            async with AsyncSessionLocal() as session:
                db_service = DatabaseService(session)
                root_folder = await db_service.create_folder(
                    name=target_dir.name,
                    path=str(target_dir)
                )

                # Process all files and subfolders
                await self._process_folder_contents(target_dir, root_folder.id, db_service)

                # Get folder structure from database
                folder_structure = await db_service.get_folder_structure(root_folder.id)
                documents = await db_service.get_all_documents(root_folder.id)

                # Calculate statistics
                total_files = len(documents)
                total_size = sum(doc.size for doc in documents)
                document_types = {}
                for doc in documents:
                    doc_type = doc.file_type.value
                    document_types[doc_type] = document_types.get(doc_type, 0) + 1

                avg_file_size = total_size / total_files if total_files > 0 else 0
                last_modified = max((doc.modified_at for doc in documents), default=datetime.now())

                return FolderAnalysis(
                    total_documents=total_files,
                    total_size=total_size,
                    document_types=document_types,
                    average_file_size=avg_file_size,
                    last_modified=last_modified,
                    documents=[self._db_to_document(doc) for doc in documents],
                    folder_structure=folder_structure
                ), total_files

        except Exception as e:
            logger.error(f"Error in analyze_folder: {str(e)}", exc_info=True)
            raise

    async def _process_folder_contents(self, folder_path: Path, folder_id: str, db_service: DatabaseService):
        """Process all contents of a folder and save to database."""
        for item in folder_path.iterdir():
            if item.name.startswith('.') or item.name in ['.DS_Store', 'Thumbs.db']:
                continue

            if item.is_file():
                try:
                    doc = await self.create_document(item, folder_path)
                    await db_service.create_document(doc, folder_id)
                except Exception as e:
                    logger.error(f"Error processing file {item}: {str(e)}")
                    continue
            elif item.is_dir():
                try:
                    subfolder = await db_service.create_folder(
                        name=item.name,
                        path=str(item),
                        parent_id=folder_id
                    )
                    await self._process_folder_contents(item, subfolder.id, db_service)
                except Exception as e:
                    logger.error(f"Error processing subfolder {item}: {str(e)}")
                    continue

    def _db_to_document(self, db_doc: DBDocument) -> Document:
        """Convert database document to Document model."""
        return Document(
            id=db_doc.id,
            filename=db_doc.filename,
            file_type=db_doc.file_type,
            size=db_doc.size,
            created_at=db_doc.created_at,
            modified_at=db_doc.modified_at,
            path=db_doc.path,
            content=db_doc.content,
            sentiment=SentimentAnalysis(
                polarity=db_doc.sentiment_polarity,
                subjectivity=db_doc.sentiment_subjectivity,
                sentiment=db_doc.sentiment_label
            ) if db_doc.sentiment_polarity is not None else None,
            metadata=db_doc.metadata,
            status=db_doc.status
        )

    async def get_document(self, document_id: str) -> Optional[Document]:
        """Get a document by ID."""
        try:
            async with AsyncSessionLocal() as session:
                db_service = DatabaseService(session)
                db_doc = await db_service.get_document(document_id)
                return self._db_to_document(db_doc) if db_doc else None
        except Exception as e:
            logger.error(f"Error getting document {document_id}: {str(e)}")
            raise

    async def delete_document(self, document_id: str) -> bool:
        """Delete a document by ID."""
        try:
            async with AsyncSessionLocal() as session:
                db_service = DatabaseService(session)
                return await db_service.delete_document(document_id)
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {str(e)}")
            raise 